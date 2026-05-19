#!/usr/bin/env python3
"""生成华为网管产品安全测试技能验证用测试文档。
聚焦三个高危模式：文件/路径相关、第三方对接、配置凭据。"""

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def h(doc, text, level=1):
    t = doc.add_heading(text, level=level)
    for r in t.runs:
        r.font.color.rgb = RGBColor(0, 51, 102)
    return t


def generate_doc1():
    """需求规格说明书 - 混合高危和低危业务功能"""
    doc = Document()

    title = doc.add_heading('华为U2000网管平台 V3.2', level=0)
    for r in title.runs:
        r.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph('需求规格说明书（节选）\n文档版本：V1.2')
    doc.add_page_break()

    h(doc, '第一章 业务功能需求', 1)

    # === 高危：文件上传功能 ===
    h(doc, '1.1 设备软件包上传', 2)
    doc.add_paragraph(
        '功能描述：管理员可通过Web界面将设备软件包（如固件、补丁文件）上传到网管服务器。'
        '上传后的软件包可用于设备升级。支持上传的文件格式包括.bin、.pat、.zip，'
        '单个文件最大500MB。上传过程中显示进度条，上传完成后系统自动校验文件完整性。'
    )

    # === 高危：文件路径/存储设置 ===
    h(doc, '1.2 备份存储路径配置', 2)
    doc.add_paragraph(
        '功能描述：管理员可以配置设备配置备份文件的存储位置。'
        '支持将备份文件存储在网管服务器本地磁盘，或通过SFTP协议存储到远程备份服务器。'
        '配置项包括：SFTP服务器IP地址、端口号、用户名、登录密码、存储目录路径。'
    )

    # === 高危：第三方系统对接 ===
    h(doc, '1.3 北向接口对接OSS系统', 2)
    doc.add_paragraph(
        '功能描述：网管平台提供北向接口供上级OSS系统对接。'
        'OSS系统可以通过RESTful接口获取全网资源数据。'
        '对接配置包括：HTTPS证书配置、Token认证参数、数据推送地址设置。'
    )

    # === 高危：系统间凭据配置 ===
    h(doc, '1.4 SNMP参数配置', 2)
    doc.add_paragraph(
        '功能描述：管理员可配置设备SNMP参数，包括SNMP版本（v1/v2c/v3）、'
        'Community字符串（读写）、SNMP v3的认证用户名和密码等。'
        '配置完成后，网管系统通过SNMP协议采集设备性能和告警数据。'
    )

    # === 低危功能（非高危） ===
    h(doc, '1.5 拓扑图展示', 3)
    doc.add_paragraph('功能描述：系统根据设备发现结果自动生成网络拓扑图，以图形化方式展示设备间的连接关系。')

    h(doc, '1.6 性能报表查看', 3)
    doc.add_paragraph('功能描述：用户可查看设备的CPU、内存、端口流量的历史性能数据报表，支持按时间范围筛选。')

    h(doc, '1.7 告警列表查询', 3)
    doc.add_paragraph('功能描述：用户可按级别、时间、设备等条件查询历史告警记录，查看告警详情。')

    path = os.path.join(OUTPUT_DIR, '华为网管平台-需求规格说明书.docx')
    doc.save(path)
    print(f'[OK] {path}')
    return path


def generate_doc2():
    """日常运维操作指南 - 混合高危和低危操作步骤"""
    doc = Document()

    title = doc.add_heading('华为U2000网管平台', level=0)
    for r in title.runs:
        r.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph('日常运维操作指南（节选）\n文档版本：V3.5')
    doc.add_page_break()

    h(doc, '第一章 系统配置操作', 1)

    # === 高危：文件路径配置 ===
    h(doc, '1.1 配置备份SFTP服务器', 2)
    doc.add_paragraph(
        '操作场景：需要将设备配置备份文件存储到远程备份服务器时，需要先配置SFTP服务器参数。'
    )
    doc.add_paragraph('操作步骤：')
    doc.add_paragraph('进入"系统设置"→"备份管理"→"存储设置"', style='List Bullet')
    doc.add_paragraph('选择存储类型为"SFTP远程服务器"', style='List Bullet')
    doc.add_paragraph('填写SFTP服务器IP地址（如192.168.1.100）和端口号（默认22）', style='List Bullet')
    doc.add_paragraph('填写登录用户名和密码', style='List Bullet')
    doc.add_paragraph('填写备份文件存储路径（如 /backup/config/）', style='List Bullet')
    doc.add_paragraph('点击"测试连接"验证SFTP连通性，测试通过后点击"保存"', style='List Bullet')

    # === 高危：文件上传 ===
    h(doc, '1.2 设备软件包上传操作', 2)
    doc.add_paragraph('操作步骤：')
    doc.add_paragraph('进入"设备管理"→"软件管理"→"上传软件包"', style='List Bullet')
    doc.add_paragraph('点击"选择文件"，选择本地的.bin或.pat文件', style='List Bullet')
    doc.add_paragraph('点击"上传"，等待上传完成，系统自动校验文件完整性', style='List Bullet')
    doc.add_paragraph('上传成功后，可在软件包列表中查看已上传的软件包', style='List Bullet')

    # === 高危：第三方对接配置 ===
    h(doc, '1.3 配置北向接口HTTPS证书', 2)
    doc.add_paragraph('操作步骤：')
    doc.add_paragraph('进入"系统设置"→"北向接口"→"证书管理"', style='List Bullet')
    doc.add_paragraph('点击"导入证书"，选择HTTPS证书文件（.pem格式）', style='List Bullet')
    doc.add_paragraph('输入证书私钥密码', style='List Bullet')
    doc.add_paragraph('点击"导入"，导入成功后重启北向接口服务使配置生效', style='List Bullet')

    # === 高危：凭据配置 ===
    h(doc, '1.4 配置设备SNMP Community', 2)
    doc.add_paragraph('操作步骤：')
    doc.add_paragraph('进入"设备管理"→"SNMP配置"', style='List Bullet')
    doc.add_paragraph('选择目标设备或设备组', style='List Bullet')
    doc.add_paragraph('设置SNMP版本（v2c/v3），配置Community字符串', style='List Bullet')
    doc.add_paragraph('如果是v3，还需要配置认证用户名、认证密码和加密密码', style='List Bullet')
    doc.add_paragraph('点击"应用"使配置生效', style='List Bullet')

    # === 高危：webhook配置 ===
    h(doc, '1.5 配置告警推送Webhook', 2)
    doc.add_paragraph('操作步骤：')
    doc.add_paragraph('进入"告警管理"→"推送配置"→"新建推送目标"', style='List Bullet')
    doc.add_paragraph('填写接收告警的URL地址（如 https://oss.example.com/webhook）', style='List Bullet')
    doc.add_paragraph('选择推送的告警级别（紧急/重要/次要/警告）', style='List Bullet')
    doc.add_paragraph('设置Token认证参数，点击"保存"', style='List Bullet')

    doc.add_page_break()
    h(doc, '第二章 日常监控操作', 1)

    # === 低危功能 ===
    h(doc, '2.1 查看设备告警', 3)
    doc.add_paragraph('操作步骤：登录系统后进入"告警管理"模块，查看当前告警列表。可按级别、设备等条件筛选。')

    h(doc, '2.2 查看性能图表', 3)
    doc.add_paragraph('操作步骤：进入"性能管理"→选择设备→选择指标→查看性能曲线。')

    h(doc, '2.3 修改个人登录密码', 3)
    doc.add_paragraph('操作步骤：点击右上角头像→"修改密码"→输入旧密码和新密码→保存。密码需满足复杂度要求。')

    h(doc, '2.4 查看操作手册', 3)
    doc.add_paragraph('操作步骤：点击界面右上角的"帮助"按钮，在线查看操作手册。')

    path = os.path.join(OUTPUT_DIR, '华为网管平台-日常运维操作指南.docx')
    doc.save(path)
    print(f'[OK] {path}')
    return path


class PDFGen(FPDF):
    def header(self):
        if self.page_no() <= 1:
            return
        self.set_font('fz', '', 7)
        self.set_text_color(100, 100, 100)
        self.set_x(self.l_margin)
        self.multi_cell(self.w - self.l_margin - self.r_margin, 8, '华为U2000网管平台 - 第三方系统对接集成指南', align='C')
        self.set_draw_color(200, 200, 200)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(3)

    def footer(self):
        self.set_y(-15)
        self.set_font('fz', '', 7)
        self.set_text_color(128, 128, 128)
        self.set_x(self.l_margin)
        self.multi_cell(self.w - self.l_margin - self.r_margin, 10, f'第 {self.page_no()} 页', align='C')

    def stitle(self, t, level=1):
        self.set_x(self.l_margin)
        sizes = {0: (16, 10, 'B'), 1: (13, 8, 'B'), 2: (10, 7, 'B'), 3: (9, 6, '')}
        s = sizes.get(level, sizes[3])
        self.set_font('fz', s[2], s[0])
        self.set_text_color([(0, 51, 102), (0, 51, 102), (51, 51, 51), (51, 51, 51)][min(level, 3)])
        self.multi_cell(self.w - self.l_margin - self.r_margin, s[1], t, align='C' if level == 0 else '')
        self.ln(4 if level <= 1 else 2)

    def txt(self, t):
        self.set_font('fz', '', 9)
        self.set_text_color(51, 51, 51)
        self.set_x(self.l_margin)
        self.multi_cell(self.w - self.l_margin - self.r_margin, 5.5, t)
        self.ln(1)

    def bul(self, t):
        self.set_font('fz', '', 9)
        x = self.l_margin + 5
        self.set_x(x)
        self.multi_cell(self.w - x - self.r_margin, 5.5, t)


def generate_pdf():
    """第三方系统对接集成指南 - 混合高危和低危"""
    pdf = PDFGen()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_font('fz', '', '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc')
    pdf.add_font('fz', 'B', '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc')
    pdf.add_page()

    pdf.stitle('华为U2000网管平台', 0)
    pdf.txt('第三方系统对接集成指南（节选）\n文档版本：V2.0\n适用范围：OSS厂商及第三方系统集成商')
    pdf.ln(3)

    pdf.stitle('第一章 对接配置操作', 1)

    # === 高危：第三方对接配置 ===
    pdf.stitle('1.1 配置北向接口参数', 2)
    pdf.txt('OSS系统对接网管前，需要在网管侧配置北向接口参数：')
    pdf.bul('配置HTTPS证书：导入CA签发的SSL证书，确保通信加密')
    pdf.bul('配置Token认证参数：设置Token有效期（默认24小时）、Token刷新策略')
    pdf.bul('配置IP白名单：设置允许访问北向接口的OSS系统IP地址')
    pdf.bul('配置推送地址：设置告警和性能数据的主动推送目标URL')

    # === 高危：文件路径配置 ===
    pdf.stitle('1.2 配置SFTP数据交换目录', 2)
    pdf.txt('OSS系统与网管系统之间通过SFTP进行数据文件交换，需要配置SFTP参数：')
    pdf.bul('SFTP服务器地址和端口（默认22）')
    pdf.bul('登录用户名和密码')
    pdf.bul('数据发送目录和接收目录路径')
    pdf.bul('文件加密方式（可选AES-256加密传输文件）')

    # === 高危：凭据配置 ===
    pdf.stitle('1.3 配置API访问凭据', 2)
    pdf.txt('第三方系统调用北向接口时需要配置API访问凭据：')
    pdf.bul('获取API Token：调用 /api/v3/auth/login 接口，传入用户名和密码获取Token')
    pdf.bul('配置Token：第三方系统在HTTP Header中携带 Authorization: Bearer {token}')
    pdf.bul('Token有效期24小时，过期后需重新获取')
    pdf.bul('支持配置多个API账号，每个账号可分配不同的数据访问权限')

    # === 高危：webhook配置 ===
    pdf.stitle('1.4 配置告警推送Webhook', 2)
    pdf.txt('通过配置Webhook实现告警信息实时推送到第三方运维平台：')
    pdf.bul('填写第三方平台的接收URL（需为HTTPS地址）')
    pdf.bul('配置推送内容：选择推送的告警级别和告警类型')
    pdf.bul('配置认证Token：第三方平台可通过Token验证推送消息来源')
    pdf.bul('支持配置主备Webhook地址，主地址推送失败自动切换到备地址')

    pdf.add_page()
    pdf.stitle('第二章 辅助功能', 1)

    pdf.stitle('2.1 查询接口版本号', 2)
    pdf.txt('调用 GET /api/v3/version 可查询当前北向接口的API版本号，无需鉴权。')

    pdf.stitle('2.2 健康检查', 2)
    pdf.txt('调用 GET /api/v3/health 可检测北向接口服务是否正常，无需鉴权，返回服务状态和时间戳。')

    pdf.stitle('2.3 对接常见问题', 2)
    pdf.txt('Token获取失败：确认账号未过期、密码正确。返回401：Token已过期需重新获取。返回429：请求频率超限。')

    path = os.path.join(OUTPUT_DIR, '华为网管平台-第三方系统对接集成指南.pdf')
    pdf.output(path)
    print(f'[OK] {path}')
    return path


if __name__ == '__main__':
    print('生成测试文档...')
    print('=' * 40)
    generate_doc1()
    generate_doc2()
    generate_pdf()
    print('=' * 40)
    print('完成！')
